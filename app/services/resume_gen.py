from openai import OpenAI
from typing import Dict, Any, Optional
from app.core.config import settings
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
import logging

logger = logging.getLogger(__name__)


class ResumeGenerator:
    def __init__(self):
        # Lazy initialization - client will be created on first use
        self._client: Optional[OpenAI] = None
    
    def _get_client(self) -> OpenAI:
        """Get or create OpenAI client with lazy initialization"""
        if self._client is None:
            try:
                # Support OpenRouter (OpenAI-compatible API)
                if settings.use_openrouter and settings.openrouter_api_key:
                    self._client = OpenAI(
                        api_key=settings.openrouter_api_key,
                        base_url="https://openrouter.ai/api/v1"
                    )
                    logger.info("Initialized OpenRouter client for resume generation")
                elif settings.openai_api_key:
                    self._client = OpenAI(api_key=settings.openai_api_key)
                    logger.info("Initialized OpenAI client for resume generation")
                else:
                    raise ValueError(
                        "No AI API key configured. Please set either OPENAI_API_KEY or "
                        "OPENROUTER_API_KEY with USE_OPENROUTER=true in your .env file."
                    )
            except Exception as e:
                logger.error(f"Failed to initialize OpenAI client: {str(e)}")
                raise ValueError(f"Failed to initialize AI client: {str(e)}") from e
        
        return self._client
    
    def _validate_api_key_available(self) -> bool:
        """Check if API key is available without initializing client"""
        if settings.use_openrouter:
            return bool(settings.openrouter_api_key)
        return bool(settings.openai_api_key)
    
    def generate_optimized_resume(
        self,
        resume_data: Dict[str, Any],
        job_description: str,
        job_requirements: list,
        missing_skills: list = None,
        original_file_path: Optional[str] = None,
        original_text: Optional[str] = None
    ) -> str:
        """Generate ATS-optimized resume content using OpenAI
        
        Args:
            resume_data: Original resume data
            job_description: Job description text
            job_requirements: List of required skills for the job
            missing_skills: List of skills missing from resume that should be added
            original_file_path: Path to original resume file (PDF/DOCX)
            original_text: Original resume text (if file path not available)
        """
        if missing_skills is None:
            missing_skills = []
        
        # Extract original text if file path provided
        if original_file_path and not original_text:
            from app.services.resume_parser import resume_parser
            try:
                original_text = resume_parser.extract_text_from_file(original_file_path)
                logger.info(f"Extracted original text from file: {len(original_text)} characters")
            except Exception as e:
                logger.warning(f"Failed to extract text from file {original_file_path}: {str(e)}")
                original_text = None
        
        # Fix "Not specified" values from original text
        if original_text:
            from app.services.resume_parser import resume_parser
            try:
                corrected_data = resume_parser.fix_not_specified_values(resume_data, original_text)
                resume_data = corrected_data
                logger.info("Fixed 'Not specified' values from original text")
            except Exception as e:
                logger.warning(f"Failed to fix 'Not specified' values: {str(e)}")
                # Continue with original resume_data if fixing fails
        
        resume_text = self._format_resume_for_prompt(resume_data)
        
        # Build missing skills section for prompt
        missing_skills_text = ""
        if missing_skills:
            missing_skills_text = f"""
MISSING SKILLS TO ADD (Required by job but not in resume):
{', '.join(missing_skills)}

IMPORTANT: These skills MUST be added to the skills array. If the user is learning these skills, 
mention them in relevant experience/projects sections to show they're being acquired.
"""
        
        # Build original text section for prompt
        original_text_section = ""
        if original_text:
            # Limit to 2000 chars to avoid token limits
            original_text_truncated = original_text[:2000]
            original_text_section = f"""
Original Resume Text (source of truth for dates, companies, roles):
{original_text_truncated}
"""
        else:
            original_text_section = "Original Resume Text: N/A (not available)"
        
        prompt = f"""You are an expert resume writer. Create a professional, ATS-optimized resume by ENHANCING the original resume below.

CRITICAL RULES - DO NOT VIOLATE:
1. PRESERVE ALL existing skills - especially those that match job requirements
2. ADD missing skills from job requirements to the skills array (see Missing Skills section below)
3. DO NOT remove any skills that are in the job requirements - they are essential
4. Use ONLY information provided in the original resume - NEVER invent facts, companies, dates, or institutions
5. Use the CORRECTED DATA below which has real dates, companies, and roles extracted from the original text
6. DO NOT use "Not specified" - use the real values from the corrected data
7. Keep ALL real company names, institution names, dates, roles, and project names EXACTLY as provided
8. Include ALL projects from the original resume - they are important
9. Only reword descriptions to be more professional and match job keywords - keep facts the same
10. Maintain reverse chronological order (most recent first)
11. If adding new skills, mention them in relevant experience/projects to show learning context

{original_text_section}

Job Description:
{job_description}

Job Requirements:
{', '.join(job_requirements)}
{missing_skills_text}
Corrected Resume Data (extracted from original text above):
{resume_text}

Instructions:
1. PRESERVE ALL existing skills - DO NOT remove any skills, especially those matching job requirements
2. ADD missing skills listed above to the skills array - these are required for the job
3. Use EVERY piece of real information from the CORRECTED DATA above (companies, institutions, dates, roles, projects)
4. Use the REAL dates, companies, and roles from the corrected data - DO NOT use "Not specified"
5. Include ALL projects from original resume - they demonstrate real work
6. Optimize descriptions to match job requirements using ATS-friendly keywords
7. Write a professional summary (2-3 sentences) based on REAL experience and projects listed
8. Keep experience in reverse chronological order (most recent first)
9. For projects: Use exact project names, enhance descriptions professionally, include technologies mentioned
10. For education: Use exact institution names and degrees - never make up universities or degrees
11. Make descriptions more impactful while staying 100% truthful
12. If adding new skills, naturally mention them in relevant experience/projects to show learning or application context

Return the optimized resume in this JSON format:
{{
    "name": "exact name from original",
    "email": "exact email from original",
    "phone": "exact phone from original or empty string",
    "summary": "2-3 sentence professional summary highlighting REAL experience and projects from original",
    "skills": ["ALL existing skills from original PLUS missing skills from job requirements - DO NOT remove any existing skills"],
    "experience": [
        {{
            "role": "exact role from corrected data (use real role, not 'Not specified')",
            "company": "exact company from corrected data (use real company, not 'Not specified')",
            "dates": "exact dates from corrected data (use real dates, not 'Not specified')",
            "location": "location from original if available, else empty",
            "description": "professionally optimized description based on original, using job keywords"
        }}
    ],
    "projects": [
        {{
            "name": "exact project name from original (include ALL projects)",
            "description": "professionally enhanced description based on original project details",
            "technologies": ["technologies mentioned in original project description"],
            "url": "url from original if available, else empty"
        }}
    ],
    "education": [
        {{
            "degree": "exact degree from corrected data (never change)",
            "institution": "exact institution from corrected data (use real institution, not 'Not specified')",
            "field": "field of study from original if available",
            "dates": "exact dates from corrected data (use real dates, not 'Not specified')",
            "gpa": "gpa from original if available, else empty"
        }}
    ]
}}"""

        # Try with primary model first, fallback to cheaper model if needed
        models_to_try = []
        if settings.use_openrouter:
            # Try cheaper models first on OpenRouter
            models_to_try = [
                ("openai/gpt-3.5-turbo", 1500),  # Cheapest option
                ("openai/gpt-4-turbo-preview", 1500),  # Better quality if credits allow
            ]
        else:
            # For OpenAI, try gpt-3.5-turbo first (cheaper), then gpt-4
            models_to_try = [
                ("gpt-3.5-turbo", 1500),  # Cheaper option
                ("gpt-4-turbo-preview", 1500),  # Better quality
            ]
        
        last_error = None
        for model, max_tokens in models_to_try:
            try:
                client = self._get_client()
                
                logger.info(f"Attempting resume generation with model: {model}, max_tokens: {max_tokens}")
                
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are an expert resume writer specializing in ATS optimization."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=max_tokens
                )
                
                content = response.choices[0].message.content
                
                # Try to extract JSON from response
                import json
                import re
                
                # Find JSON in response
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    resume_json = json.loads(json_match.group())
                    logger.info(f"Successfully generated resume using model: {model}")
                    return json.dumps(resume_json, indent=2)
                else:
                    logger.warning(f"Could not extract JSON from response, returning raw content")
                    return content
                    
            except Exception as e:
                error_str = str(e)
                last_error = e
                logger.warning(f"Failed with model {model}: {error_str}")
                
                # If it's a credit/token issue, try next model
                if "402" in error_str or "credits" in error_str.lower() or "max_tokens" in error_str.lower():
                    logger.info(f"Credit/token issue with {model}, trying next model...")
                    continue
                # If it's auth or rate limit, don't retry
                elif "401" in error_str or "unauthorized" in error_str.lower():
                    raise Exception(
                        "API authentication failed. Please check your API key configuration."
                    )
                elif "rate limit" in error_str.lower():
                    raise Exception(
                        "API rate limit exceeded. Please try again in a few moments."
                    )
        
        # If we get here, all models failed
        error_str = str(last_error) if last_error else "Unknown error"
        logger.error(f"All models failed. Last error: {error_str}")
        
        if "402" in error_str or "credits" in error_str.lower() or "max_tokens" in error_str.lower():
            raise Exception(
                "Insufficient API credits or token limit exceeded. "
                "Please add credits to your OpenRouter/OpenAI account at https://openrouter.ai/settings/credits or use a model with lower token requirements."
            )
        else:
            raise Exception(f"Error generating resume: {error_str}")
    
    def _format_resume_for_prompt(self, resume_data: Dict[str, Any]) -> str:
        """Format resume data as text for prompt with all available information"""
        text = "=== PERSONAL INFORMATION ===\n"
        text += f"Name: {resume_data.get('name', '')}\n"
        text += f"Email: {resume_data.get('email', '')}\n"
        text += f"Phone: {resume_data.get('phone', '')}\n\n"
        
        if resume_data.get('skills'):
            text += "=== SKILLS ===\n"
            text += f"{', '.join(resume_data['skills'])}\n\n"
        
        if resume_data.get('experience'):
            text += "=== WORK EXPERIENCE ===\n"
            for exp in resume_data['experience']:
                text += f"Role: {exp.get('role', '')}\n"
                text += f"Company: {exp.get('company', '')}\n"
                if exp.get('dates'):
                    text += f"Dates: {exp.get('dates', '')}\n"
                if exp.get('location'):
                    text += f"Location: {exp.get('location', '')}\n"
                text += f"Description: {exp.get('description', '')}\n\n"
        
        if resume_data.get('projects'):
            text += "=== PROJECTS ===\n"
            for proj in resume_data['projects']:
                text += f"Name: {proj.get('name', '')}\n"
                if proj.get('url'):
                    text += f"URL: {proj.get('url', '')}\n"
                text += f"Description: {proj.get('description', '')}\n"
                if proj.get('technologies'):
                    text += f"Technologies: {', '.join(proj.get('technologies', []))}\n"
                text += "\n"
        
        if resume_data.get('education'):
            text += "=== EDUCATION ===\n"
            for edu in resume_data['education']:
                text += f"Degree: {edu.get('degree', '')}\n"
                text += f"Institution: {edu.get('institution', '')}\n"
                if edu.get('field'):
                    text += f"Field: {edu.get('field', '')}\n"
                if edu.get('dates'):
                    text += f"Dates: {edu.get('dates', '')}\n"
                if edu.get('gpa'):
                    text += f"GPA: {edu.get('gpa', '')}\n"
                text += "\n"
        
        return text
    
    def export_to_pdf(self, resume_content: str, output_path: str):
        """Export resume to professional PDF format"""
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                               rightMargin=0.75*inch, leftMargin=0.75*inch,
                               topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Parse JSON content
        import json
        data = json.loads(resume_content)
        
        # Custom styles
        title_style = styles['Heading1']
        title_style.fontSize = 18
        title_style.textColor = HexColor('#1a1a1a')
        title_style.spaceAfter = 6
        
        header_style = styles['Heading2']
        header_style.fontSize = 12
        header_style.textColor = HexColor('#2c3e50')
        header_style.spaceAfter = 6
        header_style.spaceBefore = 12
        
        # Name and contact info
        name = data.get('name', '')
        email = data.get('email', '')
        phone = data.get('phone', '')
        
        story.append(Paragraph(name.upper(), title_style))
        
        # Contact information
        contact_info = []
        if email:
            contact_info.append(email)
        if phone:
            contact_info.append(phone)
        if contact_info:
            contact_para = Paragraph(' | '.join(contact_info), styles['Normal'])
            contact_para.alignment = TA_CENTER
            story.append(contact_para)
        
        story.append(Spacer(1, 0.2*inch))
        
        # Professional Summary
        if data.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
            story.append(Paragraph(data['summary'], styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        # Skills
        if data.get('skills'):
            story.append(Paragraph("TECHNICAL SKILLS", header_style))
            skills_text = ' • '.join(data['skills'])
            story.append(Paragraph(skills_text, styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        # Experience (chronological - most recent first)
        if data.get('experience'):
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", header_style))
            for exp in data['experience']:
                # Role and company
                role = exp.get('role', '')
                company = exp.get('company', '')
                dates = exp.get('dates', '')
                
                exp_header = f"<b>{role}</b>"
                if company:
                    exp_header += f" | {company}"
                if dates:
                    exp_header += f" | {dates}"
                
                story.append(Paragraph(exp_header, styles['Heading3']))
                
                # Description with bullet points
                description = exp.get('description', '')
                if description:
                    # Split into bullet points if it's a paragraph
                    if '\n' not in description and len(description) > 100:
                        # Try to create bullet points from sentences
                        sentences = description.split('. ')
                        for sent in sentences:
                            if sent.strip():
                                story.append(Paragraph(f"• {sent.strip()}.", styles['Normal']))
                    else:
                        story.append(Paragraph(description, styles['Normal']))
                
                story.append(Spacer(1, 0.1*inch))
        
        # Projects
        if data.get('projects'):
            story.append(Paragraph("PROJECTS", header_style))
            for proj in data['projects']:
                proj_name = proj.get('name', '')
                proj_desc = proj.get('description', '')
                proj_url = proj.get('url', '')
                
                if proj_name:
                    proj_header = f"<b>{proj_name}</b>"
                    if proj_url:
                        proj_header += f" | {proj_url}"
                    story.append(Paragraph(proj_header, styles['Heading3']))
                
                if proj_desc:
                    story.append(Paragraph(proj_desc, styles['Normal']))
                
                story.append(Spacer(1, 0.1*inch))
        
        # Education
        if data.get('education'):
            story.append(Paragraph("EDUCATION", header_style))
            for edu in data['education']:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                field = edu.get('field', '')
                dates = edu.get('dates', '')
                gpa = edu.get('gpa', '')
                
                edu_text = ""
                if degree:
                    edu_text += degree
                if field:
                    edu_text += f" in {field}"
                if institution:
                    edu_text += f" | {institution}"
                if dates:
                    edu_text += f" | {dates}"
                if gpa:
                    edu_text += f" | GPA: {gpa}"
                
                if edu_text:
                    story.append(Paragraph(edu_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
    
    def export_to_docx(self, resume_content: str, output_path: str):
        """Export resume to professional DOCX format"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Parse JSON content
        import json
        data = json.loads(resume_content)
        
        # Name (centered, large, bold)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(data.get('name', '').upper())
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(26, 26, 26)
        
        # Contact information (centered)
        email = data.get('email', '')
        phone = data.get('phone', '')
        contact_info = []
        if email:
            contact_info.append(email)
        if phone:
            contact_info.append(phone)
        
        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para_format = contact_para.paragraph_format
            contact_para_format.space_after = Pt(12)
        
        # Professional Summary
        if data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(data['summary'])
            doc.add_paragraph()  # Spacing
        
        # Skills
        if data.get('skills'):
            doc.add_heading('TECHNICAL SKILLS', level=1)
            skills_para = doc.add_paragraph(' • '.join(data['skills']))
            doc.add_paragraph()  # Spacing
        
        # Experience (chronological - most recent first)
        if data.get('experience'):
            doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
            for exp in data['experience']:
                role = exp.get('role', '')
                company = exp.get('company', '')
                dates = exp.get('dates', '')
                
                # Role, company, dates
                exp_header = f"{role}"
                if company:
                    exp_header += f" | {company}"
                if dates:
                    exp_header += f" | {dates}"
                
                exp_heading = doc.add_heading(exp_header, level=2)
                
                # Description
                description = exp.get('description', '')
                if description:
                    # Add as paragraph with bullet points if needed
                    if '\n' not in description and len(description) > 100:
                        sentences = description.split('. ')
                        for sent in sentences:
                            if sent.strip():
                                para = doc.add_paragraph(sent.strip() + '.', style='List Bullet')
                    else:
                        doc.add_paragraph(description)
                
                doc.add_paragraph()  # Spacing
        
        # Projects
        if data.get('projects'):
            doc.add_heading('PROJECTS', level=1)
            for proj in data['projects']:
                proj_name = proj.get('name', '')
                proj_desc = proj.get('description', '')
                proj_url = proj.get('url', '')
                
                if proj_name:
                    proj_header = proj_name
                    if proj_url:
                        proj_header += f" | {proj_url}"
                    doc.add_heading(proj_header, level=2)
                
                if proj_desc:
                    doc.add_paragraph(proj_desc)
                
                doc.add_paragraph()  # Spacing
        
        # Education
        if data.get('education'):
            doc.add_heading('EDUCATION', level=1)
            for edu in data['education']:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                field = edu.get('field', '')
                dates = edu.get('dates', '')
                gpa = edu.get('gpa', '')
                
                edu_text = ""
                if degree:
                    edu_text += degree
                if field:
                    edu_text += f" in {field}"
                if institution:
                    edu_text += f" | {institution}"
                if dates:
                    edu_text += f" | {dates}"
                if gpa:
                    edu_text += f" | GPA: {gpa}"
                
                if edu_text:
                    doc.add_paragraph(edu_text)
                doc.add_paragraph()  # Spacing
        
        doc.save(output_path)


resume_generator = ResumeGenerator()
