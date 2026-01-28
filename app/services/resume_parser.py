import pdfplumber
from docx import Document
from typing import Dict, List, Any, Optional
import re
import logging
import json

logger = logging.getLogger(__name__)

# Lazy import for spaCy to avoid startup issues
_spacy = None
def _get_spacy():
    """Lazy import spaCy to avoid compatibility issues at startup"""
    global _spacy
    if _spacy is None:
        try:
            import spacy
            _spacy = spacy
        except Exception as e:
            logger.warning(f"Failed to import spaCy: {str(e)}. Resume parsing will have limited functionality.")
            _spacy = False  # Mark as failed
    return _spacy if _spacy is not False else None


class ResumeParser:
    def __init__(self):
        # Lazy load spaCy model to avoid startup failures
        self.nlp = None
        self._spacy_loaded = False
    
    def _load_nlp(self):
        """Lazy load spaCy NLP model"""
        if not self._spacy_loaded:
            spacy_module = _get_spacy()
            if spacy_module:
                try:
                    self.nlp = spacy_module.load("en_core_web_sm")
                    logger.info("spaCy model loaded successfully")
                except OSError:
                    logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
                    self.nlp = None
                except Exception as e:
                    logger.warning(f"Failed to load spaCy model: {str(e)}")
                    self.nlp = None
            else:
                self.nlp = None
            self._spacy_loaded = True
        return self.nlp
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF resume"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        
        return self._parse_text(text)
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX resume"""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return self._parse_text(text)
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract raw text from PDF or DOCX file"""
        if file_path.endswith('.pdf'):
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text
        else:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def fix_not_specified_values(
        self, 
        parsed_data: Dict[str, Any], 
        original_text: str
    ) -> Dict[str, Any]:
        """Extract real dates, companies, and roles from original text to replace 'Not specified'"""
        corrected_data = json.loads(json.dumps(parsed_data))  # Deep copy
        
        # Find experience section in original text
        exp_section_match = re.search(
            r'(?:Experience|Work Experience|Professional Experience|Employment)[:\n](.*?)(?=\n\n(?:Education|Projects|Skills|$))',
            original_text, re.IGNORECASE | re.DOTALL
        )
        exp_section_text = exp_section_match.group(1) if exp_section_match else ""
        
        # Fix experience entries
        for i, exp in enumerate(corrected_data.get('experience', [])):
            # Try to find this experience entry in the original text
            # Look for role, company, or dates near experience section
            
            # Extract dates - handle various formats including parentheses
            if exp.get('dates') == 'Not specified' or not exp.get('dates'):
                # More flexible date pattern: handles parentheses, various separators
                date_patterns = [
                    r'\(?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|(?:\d{1,2}[/-])?\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|(?:\d{1,2}[/-])?\d{4}|Present|Current)\)?',
                    r'\((\d{4})\s*[-–]\s*(\d{4}|Present|Current)\)',
                    r'(\d{1,2}[/-]\d{4})\s*[-–]\s*(\d{1,2}[/-]\d{4}|Present|Current)',
                ]
                
                for pattern in date_patterns:
                    dates_match = re.search(pattern, exp_section_text, re.IGNORECASE)
                    if dates_match:
                        if len(dates_match.groups()) >= 2:
                            dates = f"{dates_match.group(1)} - {dates_match.group(2)}"
                            exp['dates'] = dates
                            # Remove this date from text to avoid matching again
                            exp_section_text = exp_section_text.replace(dates_match.group(0), "", 1)
                            break
            
            # Extract role - handle special characters like /, &, -
            if exp.get('role') == 'Not specified' or not exp.get('role'):
                # Pattern for roles with special characters: AI/ML & Full-Stack Developer
                role_patterns = [
                    r'^([A-Z][A-Za-z0-9\s/&\-]+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student|Specialist|Architect))',
                    r'^([A-Z][A-Za-z0-9\s/&\-]+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student))',
                ]
                
                for pattern in role_patterns:
                    role_match = re.search(pattern, exp_section_text, re.MULTILINE)
                    if role_match:
                        role = role_match.group(1).strip()
                        if role and len(role) > 3:  # Basic validation
                            exp['role'] = role
                            # Remove this role from text to avoid matching again
                            exp_section_text = exp_section_text.replace(role_match.group(0), "", 1)
                            break
            
            # Extract company - look for patterns like "at Company", "| Company", "@ Company"
            if exp.get('company') == 'Not specified' or not exp.get('company'):
                company_patterns = [
                    r'(?:at|@|\|)\s+([A-Z][A-Za-z0-9\s&.,]+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|University|College|Institute|School)?)',
                    r'([A-Z][A-Za-z0-9\s&.,]+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|University|College|Institute|School))',
                ]
                
                for pattern in company_patterns:
                    company_match = re.search(pattern, exp_section_text)
                    if company_match:
                        company = company_match.group(1).strip()
                        # Basic validation - should be reasonable length and not a common word
                        if company and 3 <= len(company) <= 100 and company.lower() not in ['the', 'and', 'or', 'at']:
                            exp['company'] = company
                            # Remove this company from text to avoid matching again
                            exp_section_text = exp_section_text.replace(company_match.group(0), "", 1)
                            break
        
        # Find education section in original text
        edu_section_match = re.search(
            r'(?:Education|Academic|Qualifications)[:\n](.*?)(?=\n\n(?:Experience|Projects|Skills|$))',
            original_text, re.IGNORECASE | re.DOTALL
        )
        edu_section_text = edu_section_match.group(1) if edu_section_match else ""
        
        # Fix education entries
        for edu in corrected_data.get('education', []):
            # Extract dates
            if edu.get('dates') == 'Not specified' or not edu.get('dates'):
                date_pattern = r'(\d{4})\s*[-–]\s*(\d{4}|Present|Current)'
                date_match = re.search(date_pattern, edu_section_text)
                if date_match:
                    dates = f"{date_match.group(1)} - {date_match.group(2)}"
                    edu['dates'] = dates
                    edu_section_text = edu_section_text.replace(date_match.group(0), "", 1)
            
            # Extract institution
            if edu.get('institution') == 'Not specified' or not edu.get('institution'):
                institution_pattern = r'([A-Z][A-Za-z0-9\s&.,]+(?:University|College|Institute|School|Academy))'
                institution_match = re.search(institution_pattern, edu_section_text)
                if institution_match:
                    institution = institution_match.group(1).strip()
                    if institution and len(institution) > 3:
                        edu['institution'] = institution
                        edu_section_text = edu_section_text.replace(institution_match.group(0), "", 1)
            
            # Extract degree
            if edu.get('degree') == 'Not specified' or not edu.get('degree'):
                degree_pattern = r'(B\.?Tech|B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?Tech|M\.?A\.?|Ph\.?D\.?|Bachelor|Master|Doctorate|Diploma)'
                degree_match = re.search(degree_pattern, edu_section_text, re.IGNORECASE)
                if degree_match:
                    degree = degree_match.group(1).strip()
                    if degree:
                        edu['degree'] = degree
                        edu_section_text = edu_section_text.replace(degree_match.group(0), "", 1)
        
        return corrected_data
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Extract structured data from resume text with improved parsing"""
        result = {
            "name": "",
            "email": "",
            "phone": "",
            "skills": [],
            "experience": [],
            "projects": [],
            "education": []
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            result["email"] = emails[0]
        
        # Extract phone
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            result["phone"] = phones[0] if isinstance(phones[0], str) else "".join(phones[0])
        
        # Extract name (first line usually, or before email)
        lines = text.split('\n')
        if lines:
            # Try to find name before email
            for i, line in enumerate(lines[:5]):
                line_clean = line.strip()
                if line_clean and '@' not in line_clean and not re.search(r'\d{3}', line_clean):
                    # Likely name if it's 2-4 words and capitalized
                    words = line_clean.split()
                    if 2 <= len(words) <= 4 and line_clean[0].isupper():
                        result["name"] = line_clean
                        break
            if not result["name"]:
                result["name"] = lines[0].strip()
        
        # Comprehensive list of technical skills to extract
        technical_skills_keywords = {
            # Programming Languages
            "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "kotlin", "swift",
            "php", "ruby", "scala", "r", "matlab", "perl", "bash", "shell", "powershell",
            # Web Frameworks
            "react", "vue", "angular", "next.js", "nextjs", "nuxt", "svelte", "ember",
            "node.js", "nodejs", "express", "nest.js", "nestjs", "django", "flask", "fastapi",
            "spring", "spring boot", "laravel", "rails", "asp.net", "dotnet", ".net",
            # Databases
            "postgresql", "postgres", "mysql", "mongodb", "redis", "elasticsearch", "cassandra",
            "dynamodb", "oracle", "sqlite", "mariadb", "neo4j", "couchdb", "firebase", "firestore",
            # Cloud & DevOps
            "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s", "terraform",
            "ansible", "jenkins", "github actions", "gitlab ci", "circleci", "travis ci",
            "prometheus", "grafana", "elk", "splunk", "new relic",
            # Tools & Technologies
            "git", "github", "gitlab", "bitbucket", "jira", "confluence", "slack", "figma",
            "webpack", "vite", "npm", "yarn", "pnpm", "maven", "gradle", "pip",
            # AI/ML
            "machine learning", "ml", "deep learning", "neural networks", "tensorflow", "pytorch",
            "keras", "scikit-learn", "sklearn", "opencv", "yolov8", "yolo", "computer vision",
            "nlp", "natural language processing", "transformer", "bert", "gpt",
            # Frontend
            "html", "css", "sass", "scss", "less", "bootstrap", "tailwind", "material-ui",
            "redux", "mobx", "zustand", "apollo", "graphql", "rest api", "rest", "soap",
            # Backend
            "microservices", "serverless", "lambda", "api gateway", "nginx", "apache",
            "rabbitmq", "kafka", "redis pub/sub", "websocket", "socket.io",
            # Mobile
            "react native", "flutter", "ionic", "xamarin", "android", "ios", "swift ui",
            # Other
            "linux", "unix", "windows", "macos", "sql", "nosql", "orm", "prisma", "sequelize",
            "typeorm", "hibernate", "jpa", "odm", "mongoose"
        }
        
        # Soft skills and non-technical phrases to filter out
        soft_skills_patterns = [
            r'communication\s+skills?',
            r'problem\s+solving',
            r'team\s+work',
            r'leadership',
            r'creativity',
            r'critical\s+thinking',
            r'time\s+management',
            r'adaptability',
            r'collaboration',
            r'been\s+technical\s+lead',
            r'additional\s+information',
            r'^[A-Z\s]+$',  # All caps (likely section headers)
            r'\.$',  # Ends with period (likely a sentence)
        ]
        
        # Extract skills from skills section
        skills_section = re.search(
            r'(?:Skills?|Technical Skills?|Technologies?|Tech Stack)[:\n](.*?)(?=\n\n|\n[A-Z][a-z]+:|Experience|Education|Projects|$)',
            text, re.IGNORECASE | re.DOTALL
        )
        extracted_skills = set()
        
        if skills_section:
            skills_text = skills_section.group(1)
            # Extract skills from comma-separated or bullet list
            skills_list = re.split(r'[,•\-\*\|]|\n', skills_text)
            for skill in skills_list:
                skill_clean = skill.strip()
                if not skill_clean or len(skill_clean) < 2:
                    continue
                
                # Filter out soft skills and sentences
                is_soft_skill = False
                for pattern in soft_skills_patterns:
                    if re.search(pattern, skill_clean, re.IGNORECASE):
                        is_soft_skill = True
                        break
                
                if is_soft_skill:
                    continue
                
                # Check if it's a known technical skill
                skill_lower = skill_clean.lower()
                for tech_skill in technical_skills_keywords:
                    if tech_skill in skill_lower or skill_lower in tech_skill:
                        # Normalize the skill name
                        if ' ' in tech_skill:
                            extracted_skills.add(tech_skill.title())
                        else:
                            extracted_skills.add(tech_skill.capitalize())
                        break
                else:
                    # If not in known list but looks technical (short, no spaces or common tech pattern)
                    if len(skill_clean) <= 30 and not re.search(r'[.!?]$', skill_clean):
                        # Check if it contains common tech patterns
                        if re.search(r'\b(js|ts|jsx|tsx|api|sdk|ide|cli|ui|ux|db|sql|nosql|ml|ai|devops|ci/cd)\b', skill_clean, re.IGNORECASE):
                            extracted_skills.add(skill_clean)
        
        # Also extract technical skills from entire resume text (experience, projects, etc.)
        text_lower = text.lower()
        for tech_skill in technical_skills_keywords:
            # Use word boundaries for better matching
            pattern = r'\b' + re.escape(tech_skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                # Normalize the skill name
                if ' ' in tech_skill:
                    extracted_skills.add(tech_skill.title())
                else:
                    extracted_skills.add(tech_skill.capitalize())
        
        # Convert to list and sort
        result["skills"] = sorted(list(extracted_skills))
        
        # Extract experience section with better parsing
        exp_section = re.search(
            r'(?:Experience|Work Experience|Professional Experience|Employment)[:\n](.*?)(?=\n\n(?:Education|Projects|Skills|$))',
            text, re.IGNORECASE | re.DOTALL
        )
        if exp_section:
            exp_text = exp_section.group(1)
            # Look for date patterns and job titles
            # Pattern: Role | Company | Dates or Role at Company (Dates)
            exp_patterns = [
                r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student))[^\n]*\n([^\n]+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})[^\n]*)',
                r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:at|@)\s+([A-Z][a-zA-Z\s]+)',
            ]
            
            # Split by common separators
            exp_entries = re.split(r'\n\n|\n(?=[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student))', exp_text)
            for entry in exp_entries:
                if not entry.strip():
                    continue
                
                # Extract dates - handle parentheses, various separators
                date_patterns = [
                    r'\(?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|(?:\d{1,2}[/-])?\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|(?:\d{1,2}[/-])?\d{4}|Present|Current)\)?',
                    r'\((\d{4})\s*[-–]\s*(\d{4}|Present|Current)\)',
                    r'(\d{1,2}[/-]\d{4})\s*[-–]\s*(\d{1,2}[/-]\d{4}|Present|Current)',
                ]
                dates = ""
                for pattern in date_patterns:
                    dates_match = re.search(pattern, entry, re.IGNORECASE)
                    if dates_match:
                        if len(dates_match.groups()) >= 2:
                            dates = f"{dates_match.group(1)} - {dates_match.group(2)}"
                            break
                
                # Extract role - handle special characters like /, &, -
                role_patterns = [
                    r'^([A-Z][A-Za-z0-9\s/&\-]+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student|Specialist|Architect))',
                    r'^([A-Z][A-Za-z0-9\s/&\-]+(?:Engineer|Developer|Manager|Analyst|Designer|Lead|Intern|Student))',
                ]
                role = ""
                for pattern in role_patterns:
                    role_match = re.search(pattern, entry, re.MULTILINE)
                    if role_match:
                        role = role_match.group(1).strip()
                        break
                
                # Extract company - look for patterns like "at Company", "| Company", "@ Company"
                company_patterns = [
                    r'(?:at|@|\|)\s+([A-Z][A-Za-z0-9\s&.,]+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|University|College|Institute|School)?)',
                    r'([A-Z][A-Za-z0-9\s&.,]+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|University|College|Institute|School))',
                ]
                company = ""
                for pattern in company_patterns:
                    company_match = re.search(pattern, entry)
                    if company_match:
                        company = company_match.group(1).strip()
                        # Basic validation - should be reasonable length
                        if 3 <= len(company) <= 100:
                            break
                        else:
                            company = ""
                
                # Extract description (rest of text)
                desc_lines = entry.split('\n')[1:] if '\n' in entry else []
                description = ' '.join([l.strip() for l in desc_lines if l.strip() and not re.search(date_pattern, l, re.IGNORECASE)])
                
                if role or company:
                    result["experience"].append({
                        "role": role or "Not specified",
                        "company": company or "Not specified",
                        "dates": dates or "Not specified",
                        "description": description
                    })
        
        # Extract projects section
        project_section = re.search(
            r'(?:Projects?|Personal Projects?)[:\n](.*?)(?=\n\n(?:Education|Experience|Skills|$))',
            text, re.IGNORECASE | re.DOTALL
        )
        if project_section:
            project_text = project_section.group(1)
            # Split projects by bullet points or new lines with titles
            project_entries = re.split(r'\n(?=[A-Z][a-zA-Z\s]+[:\-]|\d+\.|\•|\-|\*)', project_text)
            for proj in project_entries:
                if not proj.strip():
                    continue
                # Extract project name (first line or before colon)
                name_match = re.search(r'^([A-Z][a-zA-Z\s&]+(?:Project|System|App|Platform|Assistant)?)', proj, re.MULTILINE)
                name = name_match.group(1).strip() if name_match else ""
                
                # Extract description (rest)
                desc_lines = proj.split('\n')[1:] if '\n' in proj else [proj]
                description = ' '.join([l.strip() for l in desc_lines if l.strip()])
                
                # Extract URL if present
                url_match = re.search(r'(https?://[^\s]+)', proj)
                url = url_match.group(1) if url_match else ""
                
                if name or description:
                    result["projects"].append({
                        "name": name or "Project",
                        "description": description,
                        "url": url
                    })
        
        # Extract education section
        edu_section = re.search(
            r'(?:Education|Academic|Qualifications)[:\n](.*?)(?=\n\n(?:Experience|Projects|Skills|$))',
            text, re.IGNORECASE | re.DOTALL
        )
        if edu_section:
            edu_text = edu_section.group(1)
            # Look for degree patterns
            degree_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:in|of)\s+[A-Z][a-zA-Z\s]+|B\.?Tech|B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?Tech|Ph\.?D\.?)'
            institution_pattern = r'([A-Z][a-zA-Z\s&]+(?:University|College|Institute|School))'
            gpa_pattern = r'(?:GPA|CGPA|gpa|cgpa)[:\s]*([\d.]+)'
            
            edu_entries = re.split(r'\n\n|\n(?=[A-Z])', edu_text)
            for edu in edu_entries:
                if not edu.strip():
                    continue
                
                degree_match = re.search(degree_pattern, edu, re.IGNORECASE)
                degree = degree_match.group(1).strip() if degree_match else ""
                
                institution_match = re.search(institution_pattern, edu)
                institution = institution_match.group(1).strip() if institution_match else ""
                
                # Extract dates - handle various formats
                date_patterns = [
                    r'\(?(\d{4})\s*[-–]\s*(\d{4}|Present|Current)\)?',
                    r'(\d{1,2}[/-]\d{4})\s*[-–]\s*(\d{1,2}[/-]\d{4}|Present|Current)',
                ]
                dates = ""
                for pattern in date_patterns:
                    date_match = re.search(pattern, edu, re.IGNORECASE)
                    if date_match:
                        if len(date_match.groups()) >= 2:
                            dates = f"{date_match.group(1)} - {date_match.group(2)}"
                            break
                
                gpa_match = re.search(gpa_pattern, edu)
                gpa = gpa_match.group(1) if gpa_match else ""
                
                if degree or institution:
                    result["education"].append({
                        "degree": degree or "Not specified",
                        "institution": institution or "Not specified",
                        "dates": dates or "Not specified",
                        "gpa": gpa
                    })
        
        return result


# Lazy instantiation - will be created on first use to avoid spaCy import issues at startup
_resume_parser_instance = None

def get_resume_parser():
    """Get ResumeParser instance with lazy initialization"""
    global _resume_parser_instance
    if _resume_parser_instance is None:
        _resume_parser_instance = ResumeParser()
    return _resume_parser_instance

# For backward compatibility - create a simple proxy class
class ResumeParserProxy:
    """Proxy to ResumeParser that delays instantiation"""
    def __getattr__(self, name):
        return getattr(get_resume_parser(), name)

resume_parser = ResumeParserProxy()
